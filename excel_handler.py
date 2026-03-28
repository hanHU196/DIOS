import os
import sys
import json

import pandas as pd
from docx import Document
from flask import Flask, request, send_file, jsonify
from openpyxl import load_workbook
from pymongo import MongoClient
from ai_module import extract_entities
from dotenv import load_dotenv
from docxtpl import DocxTemplate
from bson import json_util
from copy import deepcopy


# Excel 工具函数
# 解析 Excel 模板，返回字段列表
def parse_excel_template(template_path):
    df = pd.read_excel(template_path, nrows=0)
    return df.columns.tolist()

# 填充 Excel 模板
def fill_excel_with_data(template_path, data, output_path):
    """填充 Excel 模板（支持多行数据）"""
    from openpyxl import load_workbook
    
    wb = load_workbook(template_path)
    ws = wb.active
    
    # 获取表头（第一行）
    headers = [cell.value for cell in ws[1]]
    print(f"📋 表头：{headers}")
    print(f"📊 数据行数：{len(data)}")
    
    # 如果只有表头，从第二行开始添加数据
    for row_idx, record in enumerate(data, start=2):
        print(f"  写入第 {row_idx} 行：{record}")
        for col_idx, header in enumerate(headers, start=1):
            value = record.get(header, '')
            ws.cell(row=row_idx, column=col_idx, value=value)
    
    wb.save(output_path)
    print(f"✅ Excel 文件已生成：{output_path}，共 {len(data)} 行数据")
#excel模板多数据源合并
def load_data_from_file(file_path):
    """从文件加载数据（支持 JSON 和 CSV）"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.json':
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if isinstance(data, dict):
                data = [data]
            return data
    elif ext == '.csv':
        df = pd.read_csv(file_path)
        return df.to_dict('records')
    else:
        raise ValueError(f"不支持的文件类型: {ext}")

def merge_data_sources(sources):
    """合并多个数据源为一个数据列表"""
    all_data = []
    db = None  # 延迟初始化 MongoDB 连接

    for src in sources:
        if isinstance(src, dict) and src.get('type') == 'mongo':
            if db is None:
                db = MongoDBHandler()
            collection = src['collection']
            query = src.get('query', {})
            data = db.query_data(collection, query)
            all_data.extend(data)
        elif isinstance(src, str):
            # 文件路径
            try:
                data = load_data_from_file(src)
                all_data.extend(data)
            except Exception as e:
                print(f"读取文件 {src} 失败: {e}")
        elif isinstance(src, list):
            all_data.extend(src)
        else:
            print(f"忽略不支持的数据源类型: {type(src)}")
    return all_data


def read_excel_directly(file_path):
    """
    直接读取 Excel 文件，返回所有数据
    不需要 AI，速度快 100 倍
    """
    try:
        # 读取 Excel 文件
        df = pd.read_excel(file_path)
        
        # 转换为字典列表
        data = df.to_dict('records')
        
        print(f"✅ 直接读取 Excel 成功，共 {len(data)} 行数据")
        return data
    except Exception as e:
        print(f"❌ 读取 Excel 失败: {e}")
        return []
# Word 工具函数
def parse_word_template(template_path):
    """
    解析 Word 模板，提取所有表格的表头字段。
    """
    doc = Document(template_path)
    all_headers = []
    for table in doc.tables:
        if len(table.rows) > 0:
            header_row = table.rows[0]
            headers = [cell.text.strip() for cell in header_row.cells]
            all_headers.append(headers)

    if not all_headers:
        print("⚠️ 未找到任何表格")
        return []

    # 检查所有表头是否一致
    first_headers = all_headers[0]
    for headers in all_headers[1:]:
        if headers != first_headers:
            print("⚠️ 警告：表格表头不一致，将返回所有表头的并集")
            # 合并所有表头（去重）
            merged = set()
            for h in all_headers:
                merged.update(h)
            return list(merged)

    return first_headers


def fill_single_table_with_records(table, records):
    """清空表格中除表头外的所有行，并用记录填充。"""
    if not records:
        return
    header_row = table.rows[0]
    headers = [cell.text.strip() for cell in header_row.cells]
    # 删除表头以下的所有行
    tbl = table._tbl
    for i in range(len(table.rows) - 1, 0, -1):
        tr = table.rows[i]._tr
        tbl.remove(tr)
    # 添加数据行
    for record in records:
        row = table.add_row()
        for col_idx, header in enumerate(headers):
            if col_idx < len(row.cells):
                value = record.get(header, '')
                row.cells[col_idx].text = str(value)


def fill_word_tables(doc, data_records, group_by='城市', extract_key_func=None):
    """遍历文档的段落和表格，按分组填充表格。"""
    """extract_entity_func: 由队友提供的函数，输入段落文本，返回实体值（如城市名）"""

    # 按 group_by 字段分组数据
    groups = {}
    for rec in data_records:
        key = rec.get(group_by)
        if key:
            groups.setdefault(key, []).append(rec)

    # 2. 遍历文档元素（段落和表格）
    elements = []
    p_iter = iter(doc.paragraphs)
    t_iter = iter(doc.tables)
    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag.endswith('p'):
            elements.append(('paragraph', next(p_iter)))
        elif tag.endswith('tbl'):
            elements.append(('table', next(t_iter)))

    current_key = None
    for elem_type, elem in elements:
        if elem_type == 'paragraph':
            if extract_key_func:
                key = extract_key_func(elem.text)
            else:
                # 使用 AI 模块提取实体（目标字段为 group_by）
                try:
                    entities = extract_entities(elem.text, [group_by])
                    key = entities.get(group_by)
                    # 如果返回的是列表，取第一个元素（假设要求单个值）
                    if isinstance(key, list):
                        key = key[0] if key else None
                except Exception as e:
                    print(f"⚠️ AI提取实体出错: {e}")
                    key = None
            if key:
                current_key = key
        else:  # table
            if current_key and current_key in groups:
                fill_single_table_with_records(elem, groups[current_key])
            else:
                print(f"⚠️ 未找到段落对应的数据（key={current_key}）")


def fill_word_with_data(template_path, data_records, output_dir, filename_prefix="output",
                        group_by='城市', extract_key_func=None):
    """填充 Word 模板（表格填充模式）。"""
    os.makedirs(output_dir, exist_ok=True)

    doc = Document(template_path)
    fill_word_tables(doc, data_records, group_by=group_by, extract_key_func=extract_key_func)
    output_path = os.path.join(output_dir, f"{filename_prefix}.docx")
    doc.save(output_path)
    print(f"✅ Word 文件已生成（表格填充）：{output_path}")

#word模板多数据源合并
def copy_table_row(table, source_row, target_row_index):
    """将 source_row 复制到 table 的 target_row_index 位置"""
    if target_row_index >= len(table.rows):
        new_row = table.add_row()
    else:
        new_row = table.insert_row(target_row_index)
    for idx, cell in enumerate(source_row.cells):
        new_cell = new_row.cells[idx]                   # 复制段落内容和样式（简单复制文本，如需保留格式可进一步处理）
        new_cell.text = cell.text                       # 注意：如果单元格内有占位符，我们会在填充阶段统一替换，这里只是复制结构
    return new_row

def replace_placeholders_in_paragraph(paragraph, data):
    """遍历 runs 替换占位符，保留样式"""
    for run in paragraph.runs:
        text = run.text
        for key, value in data.items():
            placeholder = '{{' + key + '}}'
            if placeholder in text:
                text = text.replace(placeholder, str(value))
        run.text = text
        
        
def fill_word_with_data_merged(template_path, data, output_path, table_index=0, template_row_index=1):
    """将数据列表填入 Word 模板的表格中，生成单个文档。"""
    doc = Document(template_path)
    if not doc.tables or table_index >= len(doc.tables):
        raise ValueError(f"文档中没有表格或索引 {table_index} 超出范围")

    table = doc.tables[table_index]
    if template_row_index >= len(table.rows):
        raise ValueError(f"模板行索引 {template_row_index} 超出表格行数")

    template_row = table.rows[template_row_index]
    for i, record in enumerate(data):
        new_row = copy_table_row(table, template_row, template_row_index)
        for cell in new_row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, record)
    tbl = table._tbl
    tr = template_row._tr
    tr.getparent().remove(tr)

    doc.save(output_path)
    print(f"合并填充 Word 文件已生成：{output_path}")

#使用 docxtpl 处理半结构化数据
def fill_word_with_docxtpl(template_path, context, output_path):
    """使用 docxtpl 渲染模板"""
    from docxtpl import DocxTemplate
    doc = DocxTemplate(template_path)
    doc.render(context)
    doc.save(output_path)

# MongoDB 操作类
class MongoDBHandler:
    def __init__(self, uri="mongodb://localhost:27017/", db_name="testdb"):
        self.client = MongoClient(uri)
        self.db = self.client[db_name]

    def insert_data(self, data, type_field="type"):
        """根据数据中的 type_field 字段自动选择集合插入。"""
        if isinstance(data, list):
            # 按 type 分组
            groups = {}
            for item in data:
                coll_name = item.get(type_field)
                if not coll_name:
                    raise ValueError(f"数据缺少字段 '{type_field}'，无法确定集合")
                groups.setdefault(coll_name, []).append(item)
            inserted_ids = []
            for coll_name, items in groups.items():
                coll = self.db[coll_name]
                result = coll.insert_many(items)
                inserted_ids.extend(result.inserted_ids)
            return inserted_ids
        else:
            coll_name = data.get(type_field)
            if not coll_name:
                raise ValueError(f"数据缺少字段 '{type_field}'，无法确定集合")
            coll = self.db[coll_name]
            return coll.insert_one(data).inserted_id

    def query_data(self, collection_name, query=None):
        query = query or {}
        cursor = self.db[collection_name].find(query)
        results = []
        for doc in cursor:
            doc.pop('_id', None)
            results.append(doc)
        return results

    def clear_collection(self, collection_name):
        self.db[collection_name].delete_many({})

#  测试
def run_test():
    """运行测试：Excel 和 Word 填充"""
    # 确保模板文件夹存在
    if not os.path.exists("templates"):
        os.makedirs("templates1")
        print("请将 Excel 模板放入 templates/template1.xlsx，Word 模板放入 templates/template2.docx")
        return

    # 测试数据
    test_data = [
        {"区域": "A", "村民委员会": "B", "总户数（户）": 12345, "总人数（人）": 56789,"粮食面积（亩）":159348,"粮食总产量（吨）":954625},
        {"区域": "C", "村民委员会": "D", "总户数（户）": 55555, "总人数（人）": 44444,"粮食面积（亩）":333333,"粮食总产量（吨）":111111},
    ]

    # 创建输出目录
    os.makedirs("output", exist_ok=True)

    excel_template = "templates/template1.xlsx"
    if os.path.exists(excel_template):
        fields = parse_excel_template(excel_template)
        print("Excel 模板字段：", fields)
        fill_excel_with_data(excel_template, test_data, "output/test_fill.xlsx")
    else:
        print("跳过 Excel 测试：模板文件不存在")

    word_template = "templates/template2.docx"
    if os.path.exists(word_template):
        fill_word_with_data(word_template, test_data, "output/word_output", filename_prefix="合同")
    else:
        print("跳过 Word 测试：模板文件不存在")
        
    # 1. 多数据源合并填充 Excel
    excel_template = "templates/template.xlsx"
    if os.path.exists(excel_template):
        print("\n>>> 测试多数据源合并填充 Excel")
        db = MongoDBHandler()
        mongo_test_data = [
            {"甲方": "Mongo源公司A", "乙方": "合作方X", "金额": 50000, "日期": "2025-03-01"},
            {"甲方": "Mongo源公司B", "乙方": "合作方Y", "金额": 75000, "日期": "2025-03-15"},
        ]
        db.insert_data("contracts", mongo_test_data)  # 插入到已有集合（可能已存在数据）
        # 准备 JSON 文件数据
        json_file = "output/sample_data.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump([
                {"甲方": "JSON源公司C", "乙方": "伙伴Z", "金额": 30000, "日期": "2025-03-20"}
            ], f, ensure_ascii=False)
        # 硬编码数据
        hardcoded_data = [{"甲方": "硬编码源公司D", "乙方": "测试方", "金额": 20000, "日期": "2025-03-25"}]
        # 定义数据源
        sources = [
            {"type": "mongo", "collection": "contracts", "query": {"甲方": {"$regex": "源公司"}}},  # 只取我们刚插入的示例
            hardcoded_data,
            json_file,
        ]
        # 合并数据
        merged_data = merge_data_sources(sources)
        print(f"合并后共 {len(merged_data)} 条数据")
        # 填充 Excel
        fill_excel_with_data(excel_template, merged_data, "output/merged_fill.xlsx")
        print("✅ Excel 多数据源合并填充完成：output/merged_fill.xlsx")
    else:
        print("⚠️ Excel 模板不存在，跳过 Excel 多数据源测试")

    # 2. 多数据源合并填充 Word（单文档表格）
    word_template = "templates/template.docx"
    if os.path.exists(word_template):
        print("\n>>> 测试多数据源合并填充 Word 表格（单文档）")
        if 'merged_data' not in locals():
            sources = [
                {"type": "mongo", "collection": "contracts", "query": {"甲方": {"$regex": "源公司"}}},
                [{"甲方": "硬编码源公司D", "乙方": "测试方", "金额": 20000, "日期": "2025-03-25"}],
                "output/sample_data.json",
            ]
            merged_data = merge_data_sources(sources)
        try:
            # 注意：Word 模板中需要包含一个表格，第二行为模板行（含占位符）
            fill_word_with_data_merged(
                word_template,
                merged_data,
                "output/merged_word.docx",
                table_index=0,  # 使用第一个表格
                template_row_index=1  # 第二行作为模板行
            )
            print("✅ Word 多数据源合并填充完成：output/merged_word.docx")
        except Exception as e:
            print(f"⚠️ Word 合并填充失败：{e}")
            print("  请确保 Word 模板包含一个表格，且第二行含有占位符（如 {{甲方}}）")
    else:
        print("⚠️ Word 模板不存在，跳过 Word 多数据源测试")

    #这里是MongDB测试
    db = MongoDBHandler()
    db.clear_collection("contracts")
    sample = [
        {"区域": "A", "村民委员会": "B", "总户数（户）": 12345, "总人数（人）": 56789, "粮食面积（亩）": 159348,
          "粮食总产量（吨）": 954625},
        {"区域": "C", "村民委员会": "D", "总户数（户）": 55555, "总人数（人）": 44444, "粮食面积（亩）": 333333,
         "粮食总产量（吨）": 111111},
    ]
    db.insert_data("contracts", sample)
    data_from_db = db.query_data("contracts", {})
    if os.path.exists(excel_template):
        fill_excel_with_data(excel_template, data_from_db, "output/db_fill.xlsx")
    if os.path.exists(word_template):
        fill_word_with_data(word_template, data_from_db, "output/word_db_output", filename_prefix="db_contract")
    print("测试完成！")

# Flask API
def create_app():
    app = Flask(__name__)

    @app.route('/fill/excel', methods=['POST'])
    def fill_excel():
        """Excel 填充 API"""
        f = request.files['template']
        data = request.get_json()
        temp_path = "temp_template.xlsx"
        f.save(temp_path)
        out_path = "output_excel.xlsx"
        fill_excel_with_data(temp_path, data, out_path)
        return send_file(out_path, as_attachment=True)

    @app.route('/fill/word', methods=['POST'])
    def fill_word():
        f = request.files['template']
        data = request.get_json()
        temp_path = "temp_template.docx"
        f.save(temp_path)
        # 由于可能生成多个文件，简单起见只处理第一条记录
        if data:
            first_record = [data[0]] if isinstance(data, list) else [data]
            fill_word_with_data(temp_path, first_record, "temp_word_output", "api_output")
            out_file = os.path.join("temp_word_output", "api_output_1.docx")
            return send_file(out_file, as_attachment=True)
        else:
            return "No data provided", 400

    @app.route('/api/query', methods=['POST'])
    def query_data():
        """根据传入的查询条件从 MongoDB 查询数据"""
        req_data = request.get_json()
        if not req_data:
            return jsonify({"error": "请求体必须为 JSON"}), 400

        collection_name = req_data.get('collection')
        if not collection_name:
            return jsonify({"error": "必须提供 collection 名称"}), 400

        query = req_data.get('query', {})
        limit = req_data.get('limit', 0)  # 0 表示不限制
        skip = req_data.get('skip', 0)

        # 实例化数据库处理器（使用默认连接参数，可改进为从环境变量读取）
        db_handler = MongoDBHandler()
        cursor = db_handler.db[collection_name].find(query).skip(skip)
        if limit > 0:
            cursor = cursor.limit(limit)

        results = []
        for doc in cursor:
            # 将 ObjectId 转为字符串
            doc['_id'] = str(doc['_id'])
            results.append(doc)

        # 使用 json_util 处理特殊类型（如日期），但我们已经手动转换了 _id
        return app.response_class(
            response=json.dumps(results, default=json_util.default),
            status=200,
            mimetype='application/json'
        )

    @app.route('/api/collections', methods=['GET'])
    def list_collections():
        """列出数据库中的所有集合"""
        db_handler = MongoDBHandler()
        collections = db_handler.db.list_collection_names()
        return jsonify(collections)

    @app.route('/api/insert', methods=['POST'])
    def api_insert():
        """
        请求体 JSON 格式：
        {
            "collection": "air_quality",   # 可选，如果提供则直接存入该集合
            "data": {...} 或 [...]         # 要插入的数据
        }
        如果不提供 collection，则尝试根据数据中的 "type" 字段自动判断。
        """
        req_data = request.get_json()
        if not req_data or 'data' not in req_data:
            return jsonify({"error": "请求体必须包含 data 字段"}), 400

        data = req_data['data']
        collection_name = req_data.get('collection')

        db_handler = MongoDBHandler()
        try:
            if collection_name:
                # 直接存入指定集合
                ids = db_handler.insert_data_into_collection(collection_name, data)
            else:
                # 自动判断（使用方案1的 type 字段）
                ids = db_handler.insert_data(data)  # 需要提前定义好 insert_data 自动判断逻辑
            return jsonify({"inserted_ids": [str(id) for id in (ids if isinstance(ids, list) else [ids])]})
        except Exception as e:
            return jsonify({"error": str(e)}), 400

    return app

# 命令行入口
if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--api":
        app = create_app()
        print("启动 API 服务，访问 http://127.0.0.1:5000/fill/excel 或 /fill/word")
        app.run(debug=True)
    else:
        run_test()