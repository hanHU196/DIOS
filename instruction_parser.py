# instruction_parser.py
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from deepseek_parser import DeepSeekParser
import logging

logger = logging.getLogger(__name__)

class InstructionOperator:
    def __init__(self, api_key=None):
        self.api_key = api_key
        if api_key:
            self.deepseek = DeepSeekParser(api_key)
            logger.info("✅ DeepSeek 指令解析器初始化成功")
        else:
            self.deepseek = None
            logger.warning("⚠️ 未配置 DeepSeek API Key")
    
    def execute(self, instruction, file_path):
        """执行自然语言指令"""
        if not self.deepseek:
            return {'success': False, 'error': 'DeepSeek 未配置'}
        
        # 根据文件类型选择解析方式
        if file_path.endswith('.docx'):
            file_type = "word"
        elif file_path.endswith('.xlsx'):
            file_type = "excel"
        else:
            return {'success': False, 'error': '不支持的文件类型'}
        
        # 解析指令（修复：不要传 file_type 参数）
        command = self.deepseek.parse_instruction(instruction)  # ← 这里去掉 , file_type
        logger.info(f"解析结果: {command}")
        
        if command.get('action') == 'unknown':
            return {'success': False, 'error': f'无法理解指令: {instruction}'}
        
        # 执行操作
        if file_type == "word":
            return self._execute_word(command, file_path)
        else:
            return self._execute_excel(command, file_path)
    
    def _execute_word(self, command, file_path):
        """执行 Word 操作（支持合并后的属性和数组类型的 action）"""
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document(file_path)
            action = command.get('action')
            position = command.get('position')
            target = command.get('target', 'all')
            size = command.get('font_size') or command.get('size')
            font_name = command.get('font_name') or command.get('font', '')
            font_color = command.get('font_color') or command.get('color', '')
            
            # ========== 处理 action 可能是数组的情况 ==========
            if isinstance(action, list):
                actions = action
            else:
                actions = [action] if action else []
            
            # 颜色映射
            color_map = {
                'red': RGBColor(255, 0, 0),
                'blue': RGBColor(0, 0, 255),
                'green': RGBColor(0, 255, 0),
                'black': RGBColor(0, 0, 0)
            }
            
            # 确定段落
            if position is not None and target == 'paragraph':
                if 1 <= position <= len(doc.paragraphs):
                    paragraphs = [doc.paragraphs[position - 1]]
                else:
                    return {'success': False, 'error': f'段落 {position} 超出范围'}
            else:
                paragraphs = doc.paragraphs
            
            # 应用格式到每个段落
            for para in paragraphs:
                # 字体样式应用到每个 run
                for run in para.runs:
                    # 遍历所有 action
                    for act in actions:
                        if act == 'bold':
                            run.bold = True
                        elif act == 'italic':
                            run.italic = True
                        elif act == 'underline':
                            run.underline = True
                    # 字体
                    if font_name:
                        run.font.name = font_name
                    # 字号
                    if size:
                        run.font.size = Pt(int(size))
                    # 颜色
                    if font_color and font_color.lower() in color_map:
                        run.font.color.rgb = color_map[font_color.lower()]
                
                # 对齐（段落级别）
                for act in actions:
                    if act == 'center':
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif act == 'left':
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif act == 'right':
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.save(file_path)
            logger.info(f"✅ 格式调整完成: action={actions}, font={font_name}, size={size}, color={font_color}")
            return {'success': True, 'message': '格式调整完成'}
            
        except Exception as e:
            logger.error(f"Word 操作失败: {e}")
            import traceback
            traceback.print_exc()
            return {'success': False, 'error': str(e)}
    
    def _execute_excel(self, command, file_path):
        """执行 Excel 操作"""
        try:
            wb = load_workbook(file_path)
            ws = wb.active
            action = command.get('action')
            row = command.get('row')
            col = command.get('col')
            row_start = command.get('row_start')
            row_end = command.get('row_end')
            col_start = command.get('col_start')
            col_end = command.get('col_end')
            width = command.get('width', 15)
            height = command.get('height', 20)
            font_name = command.get('font', '')
            size = command.get('size', 11)
            
            # 确定操作范围
            if row and col:
                # 单个单元格
                cells = [(row, col)]
            elif row_start and row_end:
                # 行范围
                cells = [(r, 1) for r in range(row_start, row_end + 1)]
            elif row:
                # 单行
                cells = [(row, c) for c in range(1, ws.max_column + 1)]
            elif col:
                # 单列
                cells = [(r, col) for r in range(1, ws.max_row + 1)]
            else:
                # 全部
                cells = [(r, c) for r in range(1, ws.max_row + 1) for c in range(1, ws.max_column + 1)]
            
            # 应用格式
            for r, c in cells:
                cell = ws.cell(row=r, column=c)
                
                if action == 'excel_bold':
                    cell.font = Font(bold=True)
                elif action == 'excel_italic':
                    cell.font = Font(italic=True)
                elif action == 'excel_center':
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                elif action == 'excel_left':
                    cell.alignment = Alignment(horizontal='left')
                elif action == 'excel_right':
                    cell.alignment = Alignment(horizontal='right')
                
                if font_name:
                    cell.font = Font(name=font_name, size=size)
                elif size:
                    cell.font = Font(size=size)
            
            # 列宽
            if action == 'excel_width' and col:
                col_letter = get_column_letter(col)
                ws.column_dimensions[col_letter].width = width
            
            # 行高
            if action == 'excel_height' and row:
                ws.row_dimensions[row].height = height
            
            # 合并单元格
            if action == 'excel_merge' and row_start and col_start and row_end and col_end:
                start_cell = get_column_letter(col_start) + str(row_start)
                end_cell = get_column_letter(col_end) + str(row_end)
                ws.merge_cells(f'{start_cell}:{end_cell}')
            
            wb.save(file_path)
            return {'success': True, 'message': '格式调整完成'}
            
        except Exception as e:
            logger.error(f"Excel 操作失败: {e}")
            return {'success': False, 'error': str(e)}