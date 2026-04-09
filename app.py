import os
import re
import logging
from flask import Flask, render_template, request, send_file, jsonify
from werkzeug.utils import secure_filename
from processor import DocumentProcessor
import pandas as pd
import numpy as np
import ai_module
from document_reader import DocumentReader
from db_manager import DatabaseManager
from concurrent.futures import ThreadPoolExecutor, as_completed
import io

# 初始化处理器
processor = DocumentProcessor()
reader = DocumentReader()
app = Flask(__name__)  
db = DatabaseManager()

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 配置
UPLOAD_FOLDER = 'uploads'
ALLOWED_DOC_EXTENSIONS = {'txt', 'docx', 'md', 'xlsx'}
ALLOWED_TEMPLATE_EXTENSIONS = {'xlsx', 'xls', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs('outputs', exist_ok=True)

def allowed_file(filename, allowed_set=ALLOWED_DOC_EXTENSIONS):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_set

def parse_fields_from_instruction(instruction):
    """从指令中解析字段列表"""
    import re
    
    # 匹配 "提取甲方、乙方、金额"
    match = re.search(r'提取[：:]\s*(.+)', instruction)
    if not match:
        match = re.search(r'提取\s*(.+)', instruction)
    
    if match:
        fields_str = match.group(1)
        # 按中文顿号、逗号、空格分割
        fields = re.split(r'[、，,\s]+', fields_str)
        return [f.strip() for f in fields if f.strip()]
    
    # 如果没有"提取"关键词，尝试从指令中找中文词
    words = re.split(r'[、，,\s]+', instruction)
    return [w for w in words if len(w) >= 2 and not re.search(r'[a-zA-Z0-9]', w)]

def process_word_document(doc_path, fields):
    """处理单个 Word 文档（用于并行）"""
    try:
        text = reader.read(doc_path)
        
        # 检查缓存
        cached = db.get_cached_result(doc_path, fields)
        if cached:
            return doc_path, cached, True
        
        # AI 提取
        extracted = ai_module.extract_entities_safe(text, fields)
        db.save_cache(doc_path, fields, extracted)
        return doc_path, extracted, False
        
    except Exception as e:
        logger.error(f"处理文档 {doc_path} 失败: {e}")
        return doc_path, [], False

# ---------- 导入队友模块 ----------
try:
    from document_reader import DocumentReader
    logger.info("成功导入 document_reader.DocumentReader")
    
    reader = DocumentReader()
    
    def read_documents(file_paths):
        all_text = ""
        for path in file_paths:
            logger.info(f"📖 读取文件：{path}")
            try:
                text = reader.read(path)
                logger.info(f"  读取成功，长度：{len(text)} 字符")
                all_text += text + "\n"
            except Exception as e:
                logger.error(f"读取失败：{e}")
                all_text += f"[读取失败：{path}]\n"
        return all_text
        
except ImportError as e:
    logger.error(f"导入 document_reader 失败：{e}")
    def read_documents(file_paths):
        return "模拟文档内容"

# 表格生成模块
try:
    import excel_handler
    logger.info("✅ 成功导入 excel_handler")
except ImportError as e:
    logger.error(f"❌ 导入 excel_handler 失败：{e}")
    class MockExcelHandler:
        @staticmethod
        def fill_excel_with_data(template_path, data_list, output_path):
            df = pd.DataFrame(data_list)
            df.to_excel(output_path, index=False)
            return output_path
        @staticmethod
        def parse_excel_template(template_path):
            return ["甲方", "乙方", "金额"]
        @staticmethod
        def MongoDBHandler():
            class MockMongo:
                def insert_data(self, *args, **kwargs):
                    return None
                def query_data(self, *args, **kwargs):
                    return []
            return MockMongo()
    excel_handler = MockExcelHandler()

# AI模块
try:
    from ai_module import parse_instruction, extract_entities, extract_entities_safe
    logger.info("成功导入 ai_module")
except ImportError:
    logger.warning("导入 ai_module 失败，使用模拟函数")
    def parse_instruction(instruction):
        if '填表' in instruction:
            return 'fill_form'
        return 'unknown'
    def extract_entities(text, targets):
        return {t: f'示例{t}' for t in targets}
    def extract_entities_safe(text, targets):
        return {t: f'示例{t}' for t in targets}

# ========== 页面路由（新增工作区） ==========
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')

@app.route('/workspace')
def workspace():
    """工作区页面"""
    return render_template('workspace.html')

# ========== 原有路由保持不变 ==========
@app.route('/upload', methods=['POST'])
def upload_file():
    doc_file = request.files.get('document')
    template_file = request.files.get('template')
    instruction = request.form.get('command', '').strip()

    if not doc_file or not template_file or not instruction:
        return "请上传文档和模板文件，并输入指令", 400

    if not allowed_file(doc_file.filename, ALLOWED_DOC_EXTENSIONS):
        return "文档格式不支持", 400
    if not allowed_file(template_file.filename, ALLOWED_TEMPLATE_EXTENSIONS):
        return "模板格式不支持", 400

    doc_filename = secure_filename(doc_file.filename)
    doc_path = os.path.join(app.config['UPLOAD_FOLDER'], doc_filename)
    doc_file.save(doc_path)

    template_filename = secure_filename(template_file.filename)
    if not template_filename.endswith('.xlsx'):
        template_filename = template_filename + '.xlsx'
    template_path = os.path.join(app.config['UPLOAD_FOLDER'], 'template_' + template_filename)
    template_file.save(template_path)

    output_dir = "outputs"
    os.makedirs(output_dir, exist_ok=True)

    from datetime import datetime
    output_filename = f"result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    output_path = os.path.join(output_dir, output_filename)

    try:
        result = processor.process_single(
            doc_path=doc_path,
            template_path=template_path,
            output_path=output_path,
            instruction=instruction
        )
        
        if result['success']:
            return send_file(output_path, as_attachment=True, download_name='result.xlsx')
        else:
            return f"处理失败：{result.get('error', '未知错误')}", 500
    except Exception as e:
        logger.error(f"处理异常：{e}")
        return f"处理异常：{str(e)}", 500

@app.route('/batch_process', methods=['POST'])
def batch_process():
    from datetime import datetime
    
    doc_files = request.files.getlist('documents')
    template_files = request.files.getlist('templates')
    
    if not doc_files or not template_files:
        return jsonify({'error': '请上传文档和模板文件'}), 400
    
    doc_paths = []
    for doc_file in doc_files:
        doc_filename = secure_filename(doc_file.filename)
        doc_path = os.path.join(app.config['UPLOAD_FOLDER'], doc_filename)
        doc_file.save(doc_path)
        doc_paths.append(doc_path)
    
    template_paths = []
    for template_file in template_files:
        template_filename = secure_filename(template_file.filename)
        if not template_filename.endswith('.xlsx'):
            template_filename = template_filename + '.xlsx'
        template_path = os.path.join(app.config['UPLOAD_FOLDER'], 'template_' + template_filename)
        template_file.save(template_path)
        template_paths.append(template_path)
    
    batch_dir = f"batch_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    output_dir = os.path.join('outputs', batch_dir)
    os.makedirs(output_dir, exist_ok=True)
    
    results = processor.process_batch(doc_paths, template_paths, output_dir)
    
    success_count = sum(1 for r in results if r.get('success', False))
    
    response = {
        'success': True,
        'batch_dir': batch_dir,
        'stats': {
            'total': len(results),
            'success': success_count,
            'failed': len(results) - success_count
        },
        'results': []
    }
    
    for r in results:
        result_item = {
            'template_name': os.path.basename(r['template']),
            'success': r['success']
        }
        if r['success']:
            result_item['data_count'] = r.get('data_count', 0)
            result_item['output_file'] = os.path.basename(r['output_path'])
        else:
            result_item['error'] = r.get('error', '未知错误')
        response['results'].append(result_item)
    
    return jsonify(response)

@app.route('/operate', methods=['POST'])
def operate_document():
    file = request.files.get('file')
    instruction = request.form.get('instruction', '').strip()
    
    if not file or not instruction:
        return jsonify({'error': '请上传文件并输入指令'}), 400
    
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    
    result = processor.operate_document(file_path, instruction)
    
    if result['success']:
        return send_file(file_path, as_attachment=True, download_name=f'operated_{filename}')
    else:
        return jsonify({'error': result['error']}), 400

# ========== 核心 API 接口（供前端调用） ==========
@app.route('/api/extract', methods=['POST'])
def api_extract():
    """智能文档提取（增强版）"""
    try:
        files = request.files.getlist('documents')
        command = request.form.get('command', '').strip()
        
        if not files or not command:
            return jsonify({'error': '请上传文档并输入指令'}), 400
        
        doc_paths = []
        all_text = ""
        source_files = []
        
        for file in files:
            filename = secure_filename(file.filename)
            path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(path)
            doc_paths.append(path)
            source_files.append(filename)
            
            # 保存文档信息到数据库
            text = reader.read(path)
            db.save_document(filename, path, text[:500])
            all_text += f"\n--- {filename} ---\n{text}\n"
        
        # 解析字段
        fields = parse_fields_from_instruction(command)
        if not fields:
            return jsonify({'error': '无法从指令中识别字段'}), 400
        
        # 检查缓存
        cached_result = None
        for doc_path in doc_paths:
            cached = db.get_cached_result(doc_path, fields)
            if cached:
                cached_result = cached
                break
        
        if cached_result:
            extracted = cached_result
            logger.info("📦 使用缓存结果")
        else:
            # AI 提取
            extracted = ai_module.extract_entities(all_text, fields)
            # 保存缓存
            for doc_path in doc_paths:
                db.save_cache(doc_path, fields, extracted)
        
        # 确保返回格式为列表
        if isinstance(extracted, dict):
            extracted = [extracted]
        elif not isinstance(extracted, list):
            extracted = []
        
        # 保存提取记录
        db.save_extraction(source_files, fields, extracted, command)
        
        return jsonify({
            'success': True,
            'fields': fields,
            'data': extracted,
            'from_cache': cached_result is not None
        })
        
    except Exception as e:
        logger.error(f"提取失败: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/fill', methods=['POST'])
def api_fill():
    """表格智能填写（支持并行处理）"""
    try:
        doc_files = request.files.getlist('documents')
        template_file = request.files.get('template')
        command = request.form.get('command', '').strip()
        
        if not doc_files or not template_file or not command:
            return jsonify({'error': '请上传文档、模板并输入指令'}), 400
        
        # 保存文件
        doc_paths = []
        for file in doc_files:
            filename = secure_filename(file.filename)
            path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(path)
            doc_paths.append(path)
            logger.info(f"📄 文档已保存：{filename}")
        
        template_filename = secure_filename(template_file.filename)
        template_path = os.path.join(app.config['UPLOAD_FOLDER'], 'template_' + template_filename)
        template_file.save(template_path)
        logger.info(f"📊 模板已保存：{template_filename}")
        
        # 判断模板类型
        is_word_template = template_filename.endswith('.docx')
        is_excel_template = template_filename.endswith('.xlsx') or template_filename.endswith('.xls')
        
        if not is_word_template and not is_excel_template:
            return jsonify({'error': '模板格式不支持'}), 400
        
        # ========== 特殊场景：单个 Excel 数据文件 + Word 模板 ==========
        is_excel_to_word = (len(doc_paths) == 1 and
                            doc_paths[0].endswith(('.xlsx', '.xls')) and
                            is_word_template)
        
        if is_excel_to_word:
            output_filename = f"{os.path.splitext(os.path.basename(doc_paths[0]))[0]}_填表.docx"
            output_path = os.path.join('outputs', output_filename)
            os.makedirs('outputs', exist_ok=True)
            fill_word_from_excel(template_path, doc_paths[0], output_path)
            return send_file(output_path, as_attachment=True, download_name=output_filename)
        
        # 解析模板字段
        if is_excel_template:
            from excel_handler import parse_excel_template
            fields = parse_excel_template(template_path)
        else:
            fields = parse_word_template(template_path)
        logger.info(f"📋 模板字段：{fields}")
        
        # 筛选条件
        filters = []
        try:
            from ai_module import parse_filter_conditions
            filter_result = parse_filter_conditions(command)
            filters = filter_result.get('filters', [])
        except:
            pass
        
        # ========== 并行处理 Word 文档 ==========
        all_data = []
        word_docs = []
        excel_data = []
        
        for doc_path in doc_paths:
            if doc_path.endswith('.xlsx') or doc_path.endswith('.xls'):
                df = pd.read_excel(doc_path)
                if filters:
                    df = apply_filters_to_df(df, filters)
                excel_data.extend(df.to_dict('records'))
            else:
                word_docs.append(doc_path)
        
        # 并行处理 Word 文档
        if word_docs:
            logger.info(f"🚀 并行处理 {len(word_docs)} 个 Word 文档")
            
            with ThreadPoolExecutor(max_workers=4) as executor:
                futures = {executor.submit(process_word_document, doc, fields): doc 
                           for doc in word_docs}
                
                for future in as_completed(futures):
                    doc_path, extracted, from_cache = future.result()
                    logger.info(f"   {'📦 缓存' if from_cache else '🤖 AI'} 完成: {os.path.basename(doc_path)}")
                    
                    if isinstance(extracted, list):
                        all_data.extend(extracted)
                    elif isinstance(extracted, dict):
                        all_data.append(extracted)
        
        # 添加 Excel 数据
        all_data.extend(excel_data)
        
        logger.info(f"📊 合并后共 {len(all_data)} 行数据")
        
        # ========== 生成输出文件 ==========
        # 用第一个文档名命名
        first_doc_name = os.path.splitext(os.path.basename(doc_paths[0]))[0]

        if is_excel_template:
            output_filename = f"{first_doc_name}_填表.xlsx"
            output_path = os.path.join('outputs', output_filename)
            os.makedirs('outputs', exist_ok=True)
            from excel_handler import fill_excel_with_data
            fill_excel_with_data(template_path, all_data, output_path)
        else:
            output_filename = f"{first_doc_name}_填表.docx"
            output_path = os.path.join('outputs', output_filename)
            os.makedirs('outputs', exist_ok=True)
            fill_word_with_data(template_path, all_data, output_path)
        
        # 保存历史
        try:
            db.save_fill_history(
                template=template_filename,
                doc_count=len(doc_paths),
                data_count=len(all_data),
                fields=fields,
                success=True
            )
        except:
            pass
        
        # 返回文件
        return send_file(output_path, as_attachment=True, download_name=output_filename)
        
    except Exception as e:
        logger.error(f"填表失败: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/format', methods=['POST'])
def api_format():
    """文档格式调整"""
    try:
        file = request.files.get('document')
        command = request.form.get('command', '').strip()
        
        if not file or not command:
            return jsonify({'error': '请上传文档并输入指令'}), 400
        
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        result = processor.operate_document(file_path, command)
        
        if result.get('success'):
            return send_file(file_path, as_attachment=True, download_name=f'formatted_{filename}')
        else:
            return jsonify({'error': result.get('error', '格式调整失败')}), 400
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/qa', methods=['POST'])
def api_qa():
    """智能问答接口"""
    try:
        files = request.files.getlist('documents')
        question = request.form.get('question', '').strip()
        
        if not files or not question:
            return jsonify({'error': '请上传文档并输入问题'}), 400
        
        # 1. 读取并合并所有文档内容
        all_text = ""
        for file in files:
            filename = secure_filename(file.filename)
            path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(path)
            # 使用现有的 reader 读取文档
            all_text += f"\n--- {filename} ---\n{reader.read(path)}\n"
        
        # 2. 构造 AI 提示词（截取适量字符防止超 Token）
        prompt = f"""你是一个智能文档助手。请基于以下提供的【文档内容】准确回答用户的【问题】。
要求：
1. 回答要清晰、简明扼要。
2. 如果文档中没有相关信息，请明确回答“抱歉，提供的文档中没有找到相关信息”，绝对不要编造。

【文档内容】
{all_text[:8000]} 

【问题】
{question}
"""
        
        # 3. 调用 AI 模型获取回答
        answer = ai_module.call_model(prompt, max_tokens=1500)
        
        return jsonify({'answer': answer})
        
    except Exception as e:
        logger.error(f"智能问答失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/graph', methods=['POST'])
def api_graph():
    """知识图谱生成接口"""
    try:
        files = request.files.getlist('documents')
        if not files:
            return jsonify({'error': '请上传文档'}), 400
            
        # 1. 读取合并文档
        all_text = ""
        saved_paths = []  # 记录保存的文件路径，用于最后清理
        
        for file in files:
            filename = secure_filename(file.filename)
            path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(path)
            saved_paths.append(path)
            
            try:
                # 读取文件内容
                text = reader.read(path)
                all_text += f"\n--- {filename} ---\n{text}\n"
            except Exception as read_err:
                logger.error(f"读取文件 {filename} 失败: {read_err}")
            finally:
                # 确保文件被关闭后再删除
                pass
        
        # 关闭可能打开的文件句柄
        import gc
        gc.collect()
        
        # 清理临时文件
        for path in saved_paths:
            try:
                if os.path.exists(path):
                    os.remove(path)
            except Exception as del_err:
                logger.warning(f"删除临时文件失败 {path}: {del_err}")
        
        # 如果没有提取到文本内容，返回默认数据
        if not all_text.strip():
            logger.warning("未能提取到任何文本内容，返回默认图谱")
            return jsonify(get_default_graph_data())
        
        # 2. 构造提示词
        prompt = f"""你是一个知识图谱数据抽取专家。请从以下文本中提取核心实体以及它们之间的关系。

【重要】请【严格】按照以下 JSON 格式输出，只输出纯 JSON，不要包含任何解释文字、不要包含 markdown 代码块标记。

输出格式示例：
{{
    "categories": [
        {{"name": "核心对象"}},
        {{"name": "参与方"}},
        {{"name": "属性特征"}}
    ],
    "nodes": [
        {{"id": "0", "name": "合同", "category": 0, "symbolSize": 55}},
        {{"id": "1", "name": "甲方公司", "category": 1, "symbolSize": 45}},
        {{"id": "2", "name": "乙方公司", "category": 1, "symbolSize": 45}}
    ],
    "links": [
        {{"source": "0", "target": "1", "name": "签约方"}},
        {{"source": "0", "target": "2", "name": "签约方"}}
    ]
}}

要求：
1. nodes 中的 id 必须是字符串格式的数字序列（"0", "1", "2"...）
2. links 中的 source 和 target 必须对应 nodes 中的 id
3. symbolSize 范围 20-60，核心实体用大值
4. 提取最多 15 个最重要的节点
5. 所有标点符号必须使用英文半角符号

【文本内容】
{all_text[:5000]}
"""
        
        # 3. 调用 AI 模型
        response_text = ai_module.call_model(prompt, max_tokens=2500)
        
        import json
        import re
        
        logger.info(f"AI 返回内容长度: {len(response_text)}")
        
        # 4. 清理 AI 返回内容
        clean_text = re.sub(r'```json\s*', '', response_text)
        clean_text = re.sub(r'```\s*', '', clean_text)
        clean_text = clean_text.strip()
        
        # 尝试提取 JSON 对象
        json_match = re.search(r'\{[\s\S]*\}', clean_text)
        if not json_match:
            logger.warning("无法提取 JSON，使用默认数据")
            return jsonify(get_default_graph_data())
        
        json_str = json_match.group()
        
        # 5. 修复常见的 JSON 格式问题
        json_str = re.sub(r',\s*}', '}', json_str)
        json_str = re.sub(r',\s*]', ']', json_str)
        
        try:
            graph_data = json.loads(json_str)
        except json.JSONDecodeError as je:
            logger.error(f"JSON 解析失败: {je}")
            logger.error(f"问题 JSON 片段: {json_str[:300]}")
            graph_data = get_default_graph_data()
        
        # 确保必要字段存在
        if 'nodes' not in graph_data or not graph_data['nodes']:
            graph_data['nodes'] = get_default_nodes()
        if 'links' not in graph_data:
            graph_data['links'] = get_default_links()
        if 'categories' not in graph_data:
            graph_data['categories'] = get_default_categories()
        
        # 确保 nodes 有必要的字段
        for i, node in enumerate(graph_data['nodes']):
            if 'id' not in node:
                node['id'] = str(i)
            if 'symbolSize' not in node:
                node['symbolSize'] = 40
            if 'category' not in node:
                node['category'] = 0
        
        # 确保 links 中的 source/target 是字符串
        for link in graph_data['links']:
            if 'source' in link:
                link['source'] = str(link['source'])
            if 'target' in link:
                link['target'] = str(link['target'])
        
        return jsonify(graph_data)
        
    except Exception as e:
        logger.error(f"图谱生成失败: {e}")
        import traceback
        traceback.print_exc()
        return jsonify(get_default_graph_data())

def get_default_graph_data():
    """获取默认的图谱数据"""
    return {
        "categories": get_default_categories(),
        "nodes": get_default_nodes(),
        "links": get_default_links()
    }

def get_default_categories():
    """获取默认分类"""
    return [
        {"name": "核心对象"},
        {"name": "参与方"},
        {"name": "属性特征"},
        {"name": "行为动作"}
    ]

def get_default_nodes():
    """获取默认节点"""
    return [
        {"id": "0", "name": "合同主体", "category": 0, "symbolSize": 55},
        {"id": "1", "name": "甲方公司", "category": 1, "symbolSize": 45},
        {"id": "2", "name": "乙方公司", "category": 1, "symbolSize": 45},
        {"id": "3", "name": "合同金额", "category": 2, "symbolSize": 35},
        {"id": "4", "name": "签订日期", "category": 2, "symbolSize": 35},
        {"id": "5", "name": "违约责任", "category": 3, "symbolSize": 40}
    ]

def get_default_links():
    """获取默认关系"""
    return [
        {"source": "0", "target": "1", "name": "签约方"},
        {"source": "0", "target": "2", "name": "签约方"},
        {"source": "0", "target": "3", "name": "约定金额"},
        {"source": "0", "target": "4", "name": "签署日期"},
        {"source": "0", "target": "5", "name": "包含条款"}
    ]
def get_default_graph_data():
    """获取默认的图谱数据（用于演示或出错时）"""
    return {
        "categories": [
            {"name": "核心对象"},
            {"name": "参与方"},
            {"name": "属性特征"},
            {"name": "行为动作"}
        ],
        "nodes": [
            {"id": "0", "name": "合同主体", "category": 0, "symbolSize": 55},
            {"id": "1", "name": "甲方公司", "category": 1, "symbolSize": 45},
            {"id": "2", "name": "乙方公司", "category": 1, "symbolSize": 45},
            {"id": "3", "name": "合同金额", "category": 2, "symbolSize": 35},
            {"id": "4", "name": "签订日期", "category": 2, "symbolSize": 35},
            {"id": "5", "name": "违约责任", "category": 3, "symbolSize": 40}
        ],
        "links": [
            {"source": "0", "target": "1", "name": "签约方"},
            {"source": "0", "target": "2", "name": "签约方"},
            {"source": "0", "target": "3", "name": "约定金额"},
            {"source": "0", "target": "4", "name": "签署日期"},
            {"source": "0", "target": "5", "name": "包含条款"}
        ]
    }

def get_default_categories():
    """获取默认的分类"""
    return [
        {"name": "核心对象"},
        {"name": "参与方"},
        {"name": "属性特征"},
        {"name": "行为动作"}
    ]
    

@app.route('/api/stats', methods=['GET'])
def get_stats():
    """获取系统统计信息"""
    stats = db.get_statistics()
    return jsonify(stats)

# ========== 历史文档管理 API ==========

@app.route('/api/history/list', methods=['GET'])
def get_history_list():
    """获取历史文档列表"""
    try:
        # 从数据库获取历史文档
        history_list = db.get_history_documents()
        return jsonify({
            'success': True,
            'data': history_list,
            'count': len(history_list)
        })
    except Exception as e:
        logger.error(f"获取历史列表失败: {e}")
        return jsonify({'success': False, 'error': str(e), 'data': []}), 500

@app.route('/api/history/add', methods=['POST'])
def add_to_history():
    """添加文档到历史记录（带去重）"""
    try:
        files = request.files.getlist('files')
        if not files:
            return jsonify({'success': False, 'error': '请上传文件'}), 400
        
        added = []
        skipped = []
        
        for file in files:
            filename = secure_filename(file.filename)
            file_size = file.content_length
            
            # 如果没有 content_length，需要先保存再获取大小
            if not file_size:
                temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{filename}")
                file.save(temp_path)
                file_size = os.path.getsize(temp_path)
                os.remove(temp_path)
            
            # 检查是否已存在（通过文件名和大小）
            exists = db.check_history_exists(file.filename, file_size)
            if exists:
                skipped.append(file.filename)
                continue
            
            # 保存文件到临时目录
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{filename}")
            file.save(temp_path)
            
            # 读取文件内容（用于预览）
            try:
                content_preview = reader.read(temp_path)[:200]
            except:
                content_preview = ""
            
            # 保存到数据库
            doc_id = db.save_history_document(
                filename=filename,
                original_name=file.filename,
                size=file_size,
                file_type=filename.split('.')[-1] if '.' in filename else 'unknown',
                content_preview=content_preview,
                temp_path=temp_path
            )
            
            if doc_id:
                added.append({
                    'id': str(doc_id),
                    'name': file.filename,
                    'size': file_size,
                    'type': filename.split('.')[-1] if '.' in filename else 'unknown'
                })
            
            # 清理临时文件
            try:
                os.remove(temp_path)
            except:
                pass
        
        return jsonify({
            'success': True,
            'message': f'成功添加 {len(added)} 个文档到历史' + (f'，跳过 {len(skipped)} 个重复文档' if skipped else ''),
            'data': added,
            'skipped': skipped
        })
        
    except Exception as e:
        logger.error(f"添加历史失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500
@app.route('/api/history/import', methods=['POST'])
def import_from_history():
    """从历史记录导入文档到当前工作区"""
    try:
        data = request.get_json()
        doc_ids = data.get('ids', [])
        
        if not doc_ids:
            return jsonify({'success': False, 'error': '请选择要导入的文档'}), 400
        
        imported_files = []
        for doc_id in doc_ids:
            # 从数据库获取文档
            doc = db.get_history_document_by_id(doc_id)
            if doc:
                # 创建临时文件供前端下载或使用
                temp_filename = secure_filename(doc['original_name'])
                temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"import_{temp_filename}")
                
                # 将 base64 内容写回文件
                import base64
                file_content = base64.b64decode(doc['file_content'])
                with open(temp_path, 'wb') as f:
                    f.write(file_content)
                
                imported_files.append({
                    'id': str(doc['_id']),
                    'name': doc['original_name'],
                    'size': doc['size'],
                    'temp_path': temp_path
                })
        
        return jsonify({
            'success': True,
            'message': f'成功导入 {len(imported_files)} 个文档',
            'files': imported_files
        })
        
    except Exception as e:
        logger.error(f"导入历史失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/history/delete', methods=['POST'])
def delete_history():
    """删除历史文档"""
    try:
        data = request.get_json()
        doc_ids = data.get('ids', [])
        
        if not doc_ids:
            return jsonify({'success': False, 'error': '请选择要删除的文档'}), 400
        
        deleted = db.delete_history_documents(doc_ids)
        
        return jsonify({
            'success': True,
            'message': f'成功删除 {deleted} 个文档',
            'deleted': deleted
        })
        
    except Exception as e:
        logger.error(f"删除历史失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/history/clear', methods=['POST'])
def clear_all_history():
    """清空所有历史文档"""
    try:
        db.clear_all_history()
        return jsonify({
            'success': True,
            'message': '已清空所有历史文档'
        })
        
    except Exception as e:
        logger.error(f"清空历史失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/history/search', methods=['POST'])
def search_history():
    """搜索历史文档"""
    try:
        data = request.get_json()
        keyword = data.get('keyword', '').strip()
        
        if not keyword:
            results = db.get_history_documents(limit=50)
        else:
            results = db.search_history_documents(keyword)
        
        return jsonify({
            'success': True,
            'data': results,
            'count': len(results)
        })
        
    except Exception as e:
        logger.error(f"搜索历史失败: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500





# ========== 辅助函数 ==========
def parse_word_template(template_path):
    """解析 Word 模板中的占位符"""
    from docx import Document
    import re
    doc = Document(template_path)
    fields = set()
    
    for para in doc.paragraphs:
        matches = re.findall(r'\{\{(.*?)\}\}', para.text)
        for match in matches:
            fields.add(match.strip())
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(r'\{\{(.*?)\}\}', cell.text)
                for match in matches:
                    fields.add(match.strip())
    
    return list(fields)

def fill_word_with_data(template_path, data_records, output_path):
    """填充 Word 模板"""
    from docx import Document
    doc = Document(template_path)
    
    # 替换占位符
    for record in data_records:
        for key, value in record.items():
            placeholder = f'{{{{{key}}}}}'
            for para in doc.paragraphs:
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, str(value))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
    
    doc.save(output_path)
    logger.info(f"✅ Word 文件已生成：{output_path}")

def apply_filters_to_df(df, filters):
    """应用筛选条件到 DataFrame"""
    filtered_df = df.copy()
    
    for f in filters:
        filter_type = f.get('type')
        column = f.get('column')
        
        if column not in filtered_df.columns:
            logger.warning(f"列 '{column}' 不存在，跳过筛选")
            continue
        
        if filter_type == 'date_range':
            filtered_df[column] = pd.to_datetime(filtered_df[column], errors='coerce')
            start = f.get('start')
            end = f.get('end')
            mask = (filtered_df[column] >= start) & (filtered_df[column] <= end)
            filtered_df = filtered_df[mask]
        
        elif filter_type == 'numeric':
            operator = f.get('operator')
            value = f.get('value')
            
            if operator == '>':
                filtered_df = filtered_df[filtered_df[column] > value]
            elif operator == '<':
                filtered_df = filtered_df[filtered_df[column] < value]
            elif operator == '>=':
                filtered_df = filtered_df[filtered_df[column] >= value]
            elif operator == '<=':
                filtered_df = filtered_df[filtered_df[column] <= value]
            elif operator == '==':
                filtered_df = filtered_df[filtered_df[column] == value]
    
    return filtered_df

def fill_word_from_excel(template_path, excel_path, output_path):
    """动态版：根据 Word 表格前的段落文本提取城市名，根据表头自动匹配 Excel 列名"""
    from docx import Document
    import pandas as pd
    import re

    # 1. 读取 Excel 数据
    df = pd.read_excel(excel_path)

    # 2. 打开 Word 模板
    doc = Document(template_path)
    tables = doc.tables

    # 3. 识别每个表格对应的城市
    city_table_map = {}
    table_index = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            match = re.search(r'(德州市|潍坊市|临沂市)', text)
            if match:
                city_name = match.group(1)
                next_elem = para._element.getnext()
                while next_elem is not None:
                    if next_elem.tag.endswith('tbl'):
                        if table_index < len(tables):
                            city_table_map[city_name] = tables[table_index]
                            table_index += 1
                        break
                    next_elem = next_elem.getnext()
        if table_index >= len(tables):
            break

    if not city_table_map:
        if '城市' in df.columns:
            cities = df['城市'].dropna().unique().tolist()
            for i, city in enumerate(cities):
                if i < len(tables):
                    city_table_map[city] = tables[i]

    # 4. 获取 Excel 列名
    excel_columns = df.columns.tolist()

    # 5. 填充每个表格
    for city, table in city_table_map.items():
        city_data = df[df['城市'] == city]

        header_row = table.rows[0]
        header_texts = [cell.text.strip() for cell in header_row.cells]

        col_mapping = {}
        for i, header in enumerate(header_texts):
            if header in excel_columns:
                col_mapping[i] = header
            else:
                clean_header = header.replace(' ', '').replace('（', '').replace('）', '')
                for col in excel_columns:
                    if clean_header in col.replace(' ', '').replace('（', '').replace('）', ''):
                        col_mapping[i] = col
                        break
                if i not in col_mapping:
                    col_mapping[i] = None

        existing_rows = len(table.rows) - 1
        needed_rows = len(city_data) - existing_rows
        if needed_rows > 0:
            for _ in range(needed_rows):
                table.add_row()

        for row_idx, (_, row_data) in enumerate(city_data.iterrows()):
            target_row = table.rows[row_idx + 1]
            for col_idx, excel_col in col_mapping.items():
                if col_idx < len(target_row.cells):
                    value = row_data.get(excel_col, '') if excel_col else ''
                    target_row.cells[col_idx].text = str(value)

    doc.save(output_path)
    logger.info(f"✅ Word 文件已生成（动态版本）：{output_path}")
    return output_path


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)