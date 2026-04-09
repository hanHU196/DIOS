# -*- coding: utf-8 -*-
import sys
import io
import os
import pdfplumber
from docx import Document
import pandas as pd
import markdown
import re
import zipfile
from pathlib import Path

# 强制标准输出使用UTF-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# 定义类 
class DocumentReader:
    """
    万能文档读取器
    用的时候：reader = DocumentReader()  # 创建工具箱
            content = reader.read("文件")  # 用工具箱里的read工具
    """
    
    def __init__(self):
        """
        初始化函数：创建工具箱时自动执行
        """
        # 创建一个字典：文件后缀 -> 对应的读取函数
        self.supported_formats = {
            '.txt': self._read_txt,
            '.md': self._read_md,
            '.xlsx': self._read_excel,
            '.xls': self._read_excel,
            '.docx': self._read_docx,
            '.pdf': self._read_pdf,
            '.csv': self._read_csv
        }
        
        print("✅ 文档读取器初始化成功！支持格式：", list(self.supported_formats.keys()))
    
    def read_document(self, file_path):
        """兼容旧代码的别名方法"""
        return self.read(file_path)
    
    # 统一入口函数
    def read(self, file_path):
        """
        对外统一接口
        参数：file_path 文件路径（字符串）
        返回：文件内容（字符串）
        异常：如果文件不存在或读取失败，抛出异常
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        file_name = os.path.basename(file_path)
        
        print(f"📖 正在读取：{file_name}")
        
        if ext in self.supported_formats:
            try:
                content = self.supported_formats[ext](file_path)
                print(f"✅ 读取成功！共 {len(content)} 字符")
                return content
            except Exception as e:
                raise Exception(f"读取失败：{str(e)}")
        else:
            raise ValueError(f"不支持的文件格式：{ext}，支持格式：{list(self.supported_formats.keys())}")
    
    def _read_txt(self, file_path):
        """读取TXT文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'ansi']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                print(f"  使用编码：{encoding}")
                return f"【TXT文件】{os.path.basename(file_path)}\n{'='*60}\n{content}"
            except UnicodeDecodeError:
                continue
            except Exception:
                continue
        
        with open(file_path, 'rb') as f:
            content = f.read().decode('utf-8', errors='ignore')
        print("  使用二进制模式")
        return f"【TXT文件】{os.path.basename(file_path)}\n{'='*60}\n{content}"
    
    def _read_md(self, file_path):
        """读取Markdown文件"""
        content = self._read_txt(file_path)
        try:
            html = markdown.markdown(content)
            return f"【Markdown文档】{os.path.basename(file_path)}\n{'='*60}\n【原始内容】\n{content}\n\n{'='*60}\n【HTML转换】\n{html}"
        except:
            return f"【Markdown文档】{os.path.basename(file_path)}\n{'='*60}\n{content}"
    
    def _read_excel(self, file_path):
        """读取Excel文件 - 注意：会关闭文件句柄"""
        excel_file = None
        try:
            # 使用 data_only=True 读取值而不是公式
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            all_content = []
            all_content.append(f"【Excel文件】{os.path.basename(file_path)}")
            all_content.append(f"工作表数量：{len(sheet_names)}")
            all_content.append("="*60)
            
            for sheet_idx, sheet_name in enumerate(sheet_names):
                # 读取每个工作表
                df = pd.read_excel(file_path, sheet_name=sheet_name, engine='openpyxl')
                rows, cols = df.shape
                
                all_content.append(f"\n📊 工作表 {sheet_idx+1}: {sheet_name}")
                all_content.append(f"行数: {rows}, 列数: {cols}")
                all_content.append("-"*40)
                
                with pd.option_context('display.max_rows', None, 'display.max_columns', None, 'display.width', None, 'display.max_colwidth', None):
                    df_string = df.to_string(index=True, header=True)
                    all_content.append(df_string)
                
                all_content.append("-"*40)
            
            return '\n'.join(all_content)
            
        except Exception as e:
            return f"❌ Excel读取失败：{str(e)}"
        finally:
            # 确保关闭 Excel 文件句柄
            if excel_file is not None:
                try:
                    excel_file.close()
                except:
                    pass
            # 强制垃圾回收
            import gc
            gc.collect()
    
    def _read_csv(self, file_path):
        """读取CSV文件"""
        try:
            df = pd.read_csv(file_path)
            rows, cols = df.shape
            
            result = []
            result.append(f"【CSV文件】{os.path.basename(file_path)}")
            result.append(f"总行数：{rows}，总列数：{cols}")
            result.append("="*60)
            result.append("列名：")
            result.append(", ".join([str(col) for col in df.columns]))
            result.append("-"*40)
            result.append("【完整数据】")
            
            with pd.option_context('display.max_rows', None, 'display.max_columns', None, 'display.width', None, 'display.max_colwidth', None):
                result.append(df.to_string(index=True, header=True))
            
            return '\n'.join(result)
            
        except Exception as e:
            return f"❌ CSV读取失败：{str(e)}"
    
    def _read_docx(self, file_path):
        """读取Word文档 - 注意：会关闭文档句柄"""
        doc = None
        try:
            from docx import Document
            
            # 打开Word文档
            doc = Document(file_path)
            
            result = []
            result.append(f"【Word文档】{os.path.basename(file_path)}")
            result.append("="*60)
            
            # 读取所有段落
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            result.append(f"\n【段落内容】（共{len(paragraphs)}段）")
            result.append("-"*40)
            for i, para in enumerate(paragraphs, 1):
                result.append(f"{i}. {para}")
            
            # 读取所有表格
            if doc.tables:
                result.append("\n【表格内容】")
                for table_idx, table in enumerate(doc.tables, 1):
                    result.append(f"\n表格 {table_idx}:")
                    result.append("-"*20)
                    headers = []
                    for row_idx, row in enumerate(table.rows):
                        row_cells = [cell.text.strip() for cell in row.cells]
                        
                        if not any(row_cells):
                            continue
                        
                        if not headers:
                            headers = row_cells
                            result.append(" | ".join(headers))
                        else:
                            if len(headers) == len(row_cells):
                                row_str = "，".join([f"{k}:{v}" for k, v in zip(headers, row_cells) if v])
                                result.append(row_str)
                            else:
                                result.append(" | ".join(row_cells))
            
            # 检测图片数量
            result.append("\n【图片内容】")
            result.append("-"*40)
            
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_count += 1
            
            if image_count > 0:
                result.append(f"  文档中包含 {image_count} 张图片（图片内容未识别）")
            else:
                result.append("  未找到图片")
            
            return '\n'.join(result)
            
        except Exception as e:
            import traceback
            return f"❌ Word读取失败：{str(e)}\n{traceback.format_exc()}"
        finally:
            # 关闭文档句柄
            if doc is not None:
                try:
                    # 尝试关闭文档（如果有关闭方法）
                    if hasattr(doc, 'close'):
                        doc.close()
                except:
                    pass
            # 强制垃圾回收
            import gc
            gc.collect()
    
    def _read_pdf(self, file_path):
        """读取PDF文件 - 使用 with 语句自动关闭"""
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                result = []
                result.append(f"【PDF文件】{os.path.basename(file_path)}")
                result.append(f"总页数：{total_pages}")
                result.append("="*60)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    result.append(f"\n📄 第 {page_num} 页")
                    result.append("-"*40)
                    
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        result.append(page_text)
                    else:
                        result.append("（此页无文字内容）")
                    
                    tables = page.extract_tables()
                    if tables:
                        result.append("\n表格：")
                        for table_idx, table in enumerate(tables, 1):
                            if table:
                                result.append(f"表格 {table_idx}:")
                                for row in table:
                                    row_text = ' | '.join([str(cell) if cell else '' for cell in row])
                                    if row_text.strip():
                                        result.append(f"  {row_text}")
                
                return '\n'.join(result)
            
        except Exception as e:
            return f"❌ PDF读取失败：{str(e)}"
    
    def save_to_file(self, content, output_path):
        """保存读取的内容到文件"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"✅ 内容已保存到：{output_path}")
            return True
        except Exception as e:
            print(f"❌ 保存失败：{e}")
            return False


def read_document(file_path):
    """模块级别的函数，供 app.py 直接调用"""
    reader = DocumentReader()
    return reader.read(file_path)


if __name__ == "__main__":
    print("="*60)
    print("📚 文档读取工具 - 完整输出版本")
    print("="*60)
    
    reader = DocumentReader()
    
    while True:
        print("\n" + "-"*60)
        print("请选择操作：")
        print("1. 读取单个文件（输入路径）")
        print("2. 批量读取 test_files 文件夹")
        print("3. 退出")
        print("-"*60)
        
        choice = input("请输入选项 (1/2/3): ").strip()
        
        if choice == '1':
            file_path = input("\n请输入文件路径：").strip()
            file_path = file_path.strip('"').strip("'")
            
            if not os.path.exists(file_path):
                print(f"❌ 文件不存在：{file_path}")
                continue
            
            try:
                content = reader.read(file_path)
                save = input("\n是否保存到文件？(y/n): ").strip().lower()
                if save == 'y':
                    output_path = f"output_{os.path.basename(file_path)}.txt"
                    reader.save_to_file(content, output_path)
                else:
                    print("\n" + "="*60)
                    print("📄 读取结果（完整内容）：")
                    print("="*60)
                    print(content)
                    
            except Exception as e:
                print(f"❌ 读取失败：{e}")
        
        elif choice == '2':
            test_dir = "test_files"
            
            if not os.path.exists(test_dir):
                print(f"❌ 文件夹不存在：{test_dir}")
                os.makedirs(test_dir)
                print(f"已创建文件夹：{test_dir}，请把文件放进去再试")
                continue
            
            files = os.listdir(test_dir)
            if not files:
                print(f"❌ 文件夹为空：{test_dir}")
                continue
            
            print(f"\n📊 找到 {len(files)} 个文件：")
            for i, f in enumerate(files, 1):
                ext = os.path.splitext(f)[1].lower()
                status = "✅" if ext in reader.supported_formats else "❌"
                print(f"  {i}. {status} {f}")
            
            output_dir = "output_results"
            os.makedirs(output_dir, exist_ok=True)
            print("\n开始批量读取...")
            
            for filename in files:
                ext = os.path.splitext(filename)[1].lower()
                if ext not in reader.supported_formats:
                    print(f"⏭️ 跳过不支持格式：{filename}")
                    continue
                    
                file_path = os.path.join(test_dir, filename)
                print(f"\n📄 处理：{filename}")
                
                try:
                    content = reader.read(file_path)
                    output_path = os.path.join(output_dir, f"{filename}.txt")
                    reader.save_to_file(content, output_path)
                except Exception as e:
                    print(f"❌ 处理失败：{e}")
            
            print(f"\n✅ 批量处理完成！结果保存在 {output_dir} 文件夹")
        
        elif choice == '3':
            print("\n👋 再见！")
            break
        
        else:
            print("❌ 无效选项，请重新输入")