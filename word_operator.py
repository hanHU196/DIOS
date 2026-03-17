# word_operator.py 
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

class WordOperator:
    """Word文档操作器"""
    
    def __init__(self, doc_path=None):
        """初始化，可以加载现有文档或创建新文档"""
        if doc_path and os.path.exists(doc_path):
            self.doc = Document(doc_path)
            self.path = doc_path
        else:
            self.doc = Document()
            self.path = None
        self.modified = False
    
    def apply_operations(self, operations):
        """
        应用一系列操作
        operations: [{'type': 'bold', 'target': 'title'}, ...]
        """
        results = []
        
        for op in operations:
            op_type = op.get('type')
            
            if op_type == 'bold':
                result = self._apply_bold(op.get('target', 'all'))
                results.append(f"加粗：{result}")
            
            elif op_type == 'center':
                result = self._apply_center()
                results.append(f"居中：{result}")
            
            elif op_type == 'font_size':
                result = self._apply_font_size(op.get('size', 12))
                results.append(f"字体大小：{result}")
            
            elif op_type == 'insert_table':
                result = self._insert_table(
                    op.get('rows', 3), 
                    op.get('cols', 3)
                )
                results.append(f"插入表格：{result}")
            
            elif op_type == 'insert_image':
                result = self._insert_image()
                results.append(f"插入图片：{result}")
        
        self.modified = True
        return results
    
    def _apply_bold(self, target):
        """应用加粗"""
        if target == 'title':
            # 加粗第一个段落（假设是标题）
            if self.doc.paragraphs:
                for run in self.doc.paragraphs[0].runs:
                    run.bold = True
                return "标题已加粗"
        else:
            # 加粗所有文字
            for para in self.doc.paragraphs:
                for run in para.runs:
                    run.bold = True
            return "全文已加粗"
    
    def _apply_center(self):
        """应用居中"""
        for para in self.doc.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return "已居中"
    
    def _apply_font_size(self, size):
        """设置字体大小"""
        for para in self.doc.paragraphs:
            for run in para.runs:
                run.font.size = Pt(size)
        return f"字体大小设为{size}"
    
    def _insert_table(self, rows, cols):
        """插入表格"""
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # 填充表头
        for i in range(cols):
            table.cell(0, i).text = f"列{i+1}"
        
        return f"已插入{rows}x{cols}表格"
    
    def _insert_image(self, image_path=None):
        """插入图片（需要用户提供路径）"""
        # 简化版：提示用户上传图片
        return "请上传图片文件"
    
    def save(self, output_path=None):
        """保存文档"""
        save_path = output_path or self.path or "output.docx"
        self.doc.save(save_path)
        return save_path


# 测试代码
if __name__ == "__main__":
    # 创建一个测试文档
    doc = Document()
    doc.add_paragraph("这是标题")
    doc.add_paragraph("这是正文内容，用来测试各种操作。")
    doc.save("test.docx")
    
    # 测试操作
    operator = WordOperator("test.docx")
    operations = [
        {'type': 'bold', 'target': 'title'},
        {'type': 'center'},
        {'type': 'font_size', 'size': 16},
        {'type': 'insert_table', 'rows': 3, 'cols': 4}
    ]
    
    results = operator.apply_operations(operations)
    for r in results:
        print(r)
    
    operator.save("test_output.docx")
    print("✅ 文档已保存")